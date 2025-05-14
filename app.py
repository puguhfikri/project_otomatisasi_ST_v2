from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import json
import io
import os
import csv

app = Flask(__name__)

# Fungsi buat generate nomor surat otomatis
def get_next_nomor_surat():
    now = datetime.now()
    bulan_romawi = ["I", "II", "III", "IV", "V", "VI",
                    "VII", "VIII", "IX", "X", "XI", "XII"]
    bulan = bulan_romawi[now.month - 1]
    tahun = str(now.year)

    if not os.path.exists('nomor.json'):
        nomor_data = {"last_number": 0}
    else:
        with open('nomor.json', 'r') as f:
            try:
                nomor_data = json.load(f)
            except json.JSONDecodeError:
                nomor_data = {"last_number": 0}

    nomor_data["last_number"] += 1
    nomor_urut = f"{nomor_data['last_number']:03}"

    with open('nomor.json', 'w') as f:
        json.dump(nomor_data, f)

    return f"{nomor_urut}/ST/HRD/{bulan}/{tahun}"

# Logger simpel ke .txt dan .csv
def catat_log(username, nomor_surat):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    with open("log_surat.txt", "a", encoding="utf-8") as f:
        f.write(f"[{timestamp}] {username} buat surat nomor {nomor_surat}\n")

    file_exists = os.path.isfile("log_surat.csv")
    with open("log_surat.csv", "a", encoding="utf-8", newline="") as csvfile:
        writer = csv.writer(csvfile)
        if not file_exists:
            writer.writerow(["timestamp", "username", "nomor_surat"])
        writer.writerow([timestamp, username, nomor_surat])

# Fungsi buat bikin file docx surat tugas
def buat_dokumen(data, nomor_surat):
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    doc.add_heading('SURAT TUGAS', 0)

    table_info = doc.add_table(rows=2, cols=2)
    table_info.style = 'Table Grid'
    table_info.cell(0, 0).text = "Nomor"
    table_info.cell(0, 1).text = f": {nomor_surat}"
    table_info.cell(1, 0).text = "Perihal"
    table_info.cell(1, 1).text = ": Penugasan Karyawan"

    doc.add_paragraph("\nYang bertanda tangan di bawah ini menugaskan kepada:")

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Nama'
    hdr[1].text = 'Jabatan'
    hdr[2].text = 'Unit/Divisi'

    for i in range(3):
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True

    for nama, jabatan, unit in zip(data['nama'], data['jabatan'], data['unit']):
        row = table.add_row().cells
        row[0].text = nama
        row[1].text = jabatan
        row[2].text = unit

    doc.add_paragraph(f"\nUntuk melaksanakan penugasan di {data['lokasi']}, "
                      f"mulai tanggal {data['tanggal_mulai']} sampai {data['tanggal_selesai']}.")

    doc.add_paragraph(f"\nDeskripsi Penugasan:\n{data['deskripsi']}")

    doc.add_paragraph("\nDemikian surat tugas ini dibuat untuk digunakan sebagaimana mestinya.\n")

    approval = doc.add_paragraph()
    approval.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    today = datetime.now().strftime("%d %B %Y")
    approval.add_run(f"Jakarta, {today}\n")
    approval.add_run("Disetujui oleh:\n")
    approval.add_run("Kepala Pimpinan\n\n\n")
    approval.add_run("(___________________)")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Route utama
@app.route('/', methods=['GET', 'POST'])
def form():
    if request.method == 'POST':
        data = {
            'nama': request.form.getlist('nama'),
            'jabatan': request.form.getlist('jabatan'),
            'unit': request.form.getlist('unit'),
            'lokasi': request.form['lokasi'],
            'tanggal_mulai': request.form['tanggal_mulai'],
            'tanggal_selesai': request.form['tanggal_selesai'],
            'deskripsi': request.form['deskripsi']
        }

        username = request.form.get('username', 'Anonim')
        nomor_surat = get_next_nomor_surat()
        catat_log(username, nomor_surat)

        docx_file = buat_dokumen(data, nomor_surat)
        filename = f"Surat_Tugas_{nomor_surat.replace('/', '_')}.docx"
        return send_file(docx_file, as_attachment=True, download_name=filename)

    return render_template('form.html')

if __name__ == '__main__':
    app.run(debug=True)